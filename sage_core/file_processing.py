"""
SAGE Core - File Processing
============================
Document type detection and content extraction.
"""

import os
import io
import re
import logging
import zipfile
import subprocess
import asyncio
import shutil
import tempfile
from pathlib import Path
from typing import List, Tuple

logger = logging.getLogger("SAGE-Core")


class PDFProcessingError(Exception):
    """Exception raised when PDF processing fails."""
    pass

# olmocr configuration from environment
OLMOCR_SERVER = os.getenv("OLMOCR_SERVER", "")
OLMOCR_API_KEY = os.getenv("OLMOCR_API_KEY", "")
OLMOCR_MODEL = os.getenv("OLMOCR_MODEL", "allenai/olmOCR-2-7B-1025-FP8")
PDF_TIMEOUT = int(os.getenv("PDF_TIMEOUT", "600"))  # 10 minute default

# Optional imports
try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import openpyxl
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False


def detect_file_type(filename: str, content: bytes) -> str:
    """Detect file type based on extension and content."""
    ext = Path(filename).suffix.lower()

    if ext in ['.md', '.markdown']:
        return 'markdown'
    elif ext in ['.html', '.htm']:
        return 'html'
    elif ext == '.txt':
        return 'text'
    elif ext == '.pdf':
        return 'pdf'
    elif ext == '.zip':
        return 'zip'
    elif ext == '.docx':
        return 'docx'
    elif ext in ['.xlsx', '.xls']:
        return 'excel'
    elif ext in ['.rst', '.asciidoc', '.adoc']:
        return 'text'  # Treat as plain text
    else:
        # Try to detect from content
        try:
            text = content.decode('utf-8', errors='ignore')[:1000]
            if text.strip().startswith('<!DOCTYPE') or '<html' in text.lower():
                return 'html'
            elif text.startswith('---\n') or re.search(r'^#\s+\w', text, re.MULTILINE):
                return 'markdown'
        except Exception:
            pass
        return 'text'


def convert_html_to_markdown(html_content: str) -> str:
    """Convert HTML to clean Markdown."""
    from markdownify import markdownify as md
    from bs4 import BeautifulSoup

    # Clean up with BeautifulSoup first
    soup = BeautifulSoup(html_content, 'html.parser')

    # Remove script and style elements
    for script in soup(["script", "style", "nav", "footer", "header"]):
        script.decompose()

    # Convert to markdown
    markdown = md(
        str(soup),
        heading_style="atx",
        code_language_callback=lambda el: el.get('data-language') or el.get('class', [''])[0] if el.get('class') else ''
    )

    return markdown.strip()


def extract_pdf_text(pdf_content: bytes) -> str:
    """
    Extract text from PDF file using olmocr for layout preservation.
    
    DEPRECATED: This synchronous version blocks the worker process.
    Use extract_pdf_text_async() for better performance and error handling.
    
    Returns empty string on error for backward compatibility.
    """
    try:
        # Write PDF to temp file
        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
            tmp.write(pdf_content)
            tmp_path = tmp.name

        # Create workspace for olmocr output
        workspace = tempfile.mkdtemp(prefix="olmocr_")

        logger.info("Converting PDF with olmocr (this may take a while)...")

        # Build command
        cmd = [
            "python", "-m", "olmocr.pipeline", workspace,
            "--markdown", "--pdfs", tmp_path,
            "--model", OLMOCR_MODEL
        ]

        # Add server/API configuration if provided
        if OLMOCR_SERVER:
            cmd.extend(["--server", OLMOCR_SERVER])
        if OLMOCR_API_KEY:
            cmd.extend(["--api_key", OLMOCR_API_KEY])

        # Run olmocr pipeline with timeout
        result = subprocess.run(
            cmd,
            capture_output=True,
            text=True,
            timeout=PDF_TIMEOUT
        )

        if result.returncode != 0:
            logger.error(f"olmocr failed: {result.stderr}")
            return ""

        # Read generated markdown
        pdf_stem = Path(tmp_path).stem
        md_file = Path(workspace) / "markdown" / f"{pdf_stem}.md"

        if md_file.exists():
            markdown = md_file.read_text()
            logger.info(f"PDF conversion complete: {len(markdown)} chars")
        else:
            logger.warning(f"olmocr did not produce markdown output for {tmp_path}")
            markdown = ""

        # Clean up
        os.remove(tmp_path)
        shutil.rmtree(workspace, ignore_errors=True)

        return markdown
    except subprocess.TimeoutExpired:
        logger.error("olmocr timed out processing PDF")
        return ""
    except Exception as e:
        logger.error(f"Error extracting PDF with olmocr: {e}")
        return ""


async def extract_pdf_text_async(pdf_content: bytes) -> str:
    """
    Extract text from PDF file using olmocr for layout preservation (async version).
    
    This non-blocking version uses asyncio.create_subprocess_exec() to avoid blocking
    worker processes during long PDF processing operations.
    
    Args:
        pdf_content: Raw PDF file bytes
        
    Returns:
        Extracted markdown text
        
    Raises:
        PDFProcessingError: If PDF processing fails or times out
    """
    tmp_path = None
    workspace = None
    
    try:
        # Write PDF to temp file
        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
            tmp.write(pdf_content)
            tmp_path = tmp.name

        # Create workspace for olmocr output
        workspace = tempfile.mkdtemp(prefix="olmocr_")

        logger.info("Converting PDF with olmocr (async, non-blocking)...")

        # Build command
        cmd = [
            "python", "-m", "olmocr.pipeline", workspace,
            "--markdown", "--pdfs", tmp_path,
            "--model", OLMOCR_MODEL
        ]

        # Add server/API configuration if provided
        if OLMOCR_SERVER:
            cmd.extend(["--server", OLMOCR_SERVER])
        if OLMOCR_API_KEY:
            cmd.extend(["--api_key", OLMOCR_API_KEY])

        # Run olmocr pipeline asynchronously with timeout
        process = await asyncio.create_subprocess_exec(
            *cmd,
            stdout=asyncio.subprocess.PIPE,
            stderr=asyncio.subprocess.PIPE
        )

        try:
            stdout, stderr = await asyncio.wait_for(
                process.communicate(),
                timeout=PDF_TIMEOUT
            )
        except asyncio.TimeoutError:
            # Kill the process if it's still running
            try:
                process.kill()
                # Wait for process to be killed, but don't hang
                try:
                    await asyncio.wait_for(process.wait(), timeout=5)
                except asyncio.TimeoutError:
                    pass  # Process didn't terminate, but we tried
            except Exception:
                pass
            raise PDFProcessingError(
                f"PDF processing timed out after {PDF_TIMEOUT} seconds (10 minutes). "
                "The PDF may be too large or complex."
            )

        # Check return code
        if process.returncode != 0:
            stderr_text = stderr.decode('utf-8', errors='ignore') if stderr else "Unknown error"
            logger.error(f"olmocr failed: {stderr_text}")
            raise PDFProcessingError(
                f"olmocr failed with exit code {process.returncode}: {stderr_text}"
            )

        # Read generated markdown
        pdf_stem = Path(tmp_path).stem
        md_file = Path(workspace) / "markdown" / f"{pdf_stem}.md"

        if md_file.exists():
            markdown = md_file.read_text()
            logger.info(f"PDF conversion complete: {len(markdown)} chars")
            return markdown
        else:
            logger.warning(f"olmocr did not produce markdown output for {tmp_path}")
            raise PDFProcessingError(
                f"olmocr completed but did not produce expected markdown output. "
                f"Check that the PDF is valid and olmocr is configured correctly."
            )

    except PDFProcessingError:
        # Re-raise our custom exceptions
        raise
    except Exception as e:
        logger.error(f"Error extracting PDF with olmocr: {e}")
        raise PDFProcessingError(f"Unexpected error during PDF processing: {str(e)}")
    finally:
        # Always clean up temporary files
        if tmp_path and os.path.exists(tmp_path):
            try:
                os.remove(tmp_path)
            except Exception as e:
                logger.warning(f"Failed to remove temp file {tmp_path}: {e}")
        
        if workspace and os.path.exists(workspace):
            try:
                shutil.rmtree(workspace, ignore_errors=True)
            except Exception as e:
                logger.warning(f"Failed to remove workspace {workspace}: {e}")


def extract_docx_text(docx_content: bytes) -> str:
    """Extract text from DOCX file."""
    if not DOCX_AVAILABLE:
        logger.warning("python-docx not available, skipping DOCX extraction")
        return ""

    try:
        doc = DocxDocument(io.BytesIO(docx_content))
        text_parts = []

        for para in doc.paragraphs:
            if para.text.strip():
                # Check if it's a heading
                if para.style.name.startswith('Heading'):
                    level = para.style.name.replace('Heading ', '')
                    try:
                        level = int(level)
                        text_parts.append(f"{'#' * level} {para.text}")
                    except Exception:
                        text_parts.append(f"## {para.text}")
                else:
                    text_parts.append(para.text)

        # Also extract text from tables
        for table in doc.tables:
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(' | '.join(cells))
            if rows:
                text_parts.append('\n'.join(rows))

        return '\n\n'.join(text_parts)
    except Exception as e:
        logger.error(f"Error extracting DOCX: {e}")
        return ""


def extract_excel_text(excel_content: bytes) -> str:
    """Extract text from Excel file."""
    if not EXCEL_AVAILABLE:
        logger.warning("openpyxl not available, skipping Excel extraction")
        return ""

    try:
        wb = openpyxl.load_workbook(io.BytesIO(excel_content), data_only=True)
        text_parts = []

        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            text_parts.append(f"## {sheet_name}")

            rows = []
            for row in sheet.iter_rows(values_only=True):
                # Convert None to empty string and join
                cells = [str(cell) if cell is not None else '' for cell in row]
                if any(cells):  # Only add non-empty rows
                    rows.append(' | '.join(cells))

            if rows:
                text_parts.append('\n'.join(rows))

        return '\n\n'.join(text_parts)
    except Exception as e:
        logger.error(f"Error extracting Excel: {e}")
        return ""


def extract_title_from_content(content: str, filename: str) -> str:
    """Extract title from content or use filename."""
    import yaml

    # Try YAML frontmatter
    if content.startswith('---'):
        try:
            end = content.find('---', 3)
            if end > 0:
                frontmatter = yaml.safe_load(content[3:end])
                if isinstance(frontmatter, dict) and 'title' in frontmatter:
                    return frontmatter['title']
        except Exception:
            pass

    # Try to find markdown header
    match = re.search(r'^#\s+(.+)$', content, re.MULTILINE)
    if match:
        return match.group(1).strip()

    # Use filename without extension
    return Path(filename).stem.replace('_', ' ').replace('-', ' ').title()


def process_file(
    content: bytes,
    filename: str,
    library: str,
    version: str = "latest"
) -> str:
    """
    Process a single file and return its markdown content.
    
    DEPRECATED: This synchronous version blocks for PDF processing.
    Use process_file_async() for better performance with PDFs.
    """
    file_type = detect_file_type(filename, content)

    if file_type == 'markdown':
        return content.decode('utf-8', errors='ignore')
    elif file_type == 'html':
        html = content.decode('utf-8', errors='ignore')
        return convert_html_to_markdown(html)
    elif file_type == 'text':
        return content.decode('utf-8', errors='ignore')
    elif file_type == 'pdf':
        return extract_pdf_text(content)
    elif file_type == 'docx':
        return extract_docx_text(content)
    elif file_type == 'excel':
        return extract_excel_text(content)
    else:
        return content.decode('utf-8', errors='ignore')


async def process_file_async(
    content: bytes,
    filename: str,
    library: str,
    version: str = "latest"
) -> str:
    """
    Process a single file and return its markdown content (async version).
    
    This version uses async PDF extraction to avoid blocking during long operations.
    
    Args:
        content: Raw file bytes
        filename: Original filename
        library: Library name
        version: Version string
        
    Returns:
        Markdown content
        
    Raises:
        PDFProcessingError: If PDF processing fails
    """
    file_type = detect_file_type(filename, content)

    if file_type == 'markdown':
        return content.decode('utf-8', errors='ignore')
    elif file_type == 'html':
        html = content.decode('utf-8', errors='ignore')
        return convert_html_to_markdown(html)
    elif file_type == 'text':
        return content.decode('utf-8', errors='ignore')
    elif file_type == 'pdf':
        # Use async version for PDFs
        return await extract_pdf_text_async(content)
    elif file_type == 'docx':
        return extract_docx_text(content)
    elif file_type == 'excel':
        return extract_excel_text(content)
    else:
        return content.decode('utf-8', errors='ignore')


def process_zip(
    zip_content: bytes,
    library: str,
    version: str = "latest",
    max_entries: int = 500
) -> List[Tuple[str, str]]:
    """
    Process a ZIP file and return list of (filename, markdown_content) tuples.
    
    DEPRECATED: This synchronous version blocks for PDF processing.
    Use process_zip_async() for better performance with PDFs.
    
    Args:
        zip_content: Raw ZIP file bytes
        library: Library name
        version: Version string
        max_entries: Maximum number of files to process (safety limit)
    """
    files = []

    try:
        with zipfile.ZipFile(io.BytesIO(zip_content), 'r') as zf:
            entries = zf.namelist()

            # Enforce entry limit
            if len(entries) > max_entries:
                logger.warning(f"ZIP has {len(entries)} entries, limiting to {max_entries}")
                entries = entries[:max_entries]

            for name in entries:
                # Skip directories and hidden files
                if name.endswith('/') or '/.' in name or name.startswith('.'):
                    continue

                # Skip non-document files
                ext = Path(name).suffix.lower()
                if ext not in ['.md', '.markdown', '.html', '.htm', '.txt', '.pdf', '.rst', '.docx', '.xlsx', '.xls']:
                    continue

                try:
                    content = zf.read(name)
                    markdown = process_file(content, name, library, version)
                    if markdown.strip():
                        files.append((name, markdown))
                except Exception as e:
                    logger.warning(f"Error processing {name} in ZIP: {e}")
    except Exception as e:
        logger.error(f"Error reading ZIP file: {e}")

    return files


async def process_zip_async(
    zip_content: bytes,
    library: str,
    version: str = "latest",
    max_entries: int = 500
) -> List[Tuple[str, str]]:
    """
    Process a ZIP file and return list of (filename, markdown_content) tuples (async version).
    
    This version uses async file processing to avoid blocking during PDF extraction.
    
    Args:
        zip_content: Raw ZIP file bytes
        library: Library name
        version: Version string
        max_entries: Maximum number of files to process (safety limit)
        
    Returns:
        List of (filename, markdown_content) tuples
    """
    files = []

    try:
        with zipfile.ZipFile(io.BytesIO(zip_content), 'r') as zf:
            entries = zf.namelist()

            # Enforce entry limit
            if len(entries) > max_entries:
                logger.warning(f"ZIP has {len(entries)} entries, limiting to {max_entries}")
                entries = entries[:max_entries]

            for name in entries:
                # Skip directories and hidden files
                if name.endswith('/') or '/.' in name or name.startswith('.'):
                    continue

                # Skip non-document files
                ext = Path(name).suffix.lower()
                if ext not in ['.md', '.markdown', '.html', '.htm', '.txt', '.pdf', '.rst', '.docx', '.xlsx', '.xls']:
                    continue

                try:
                    content = zf.read(name)
                    # Use async version for file processing
                    markdown = await process_file_async(content, name, library, version)
                    if markdown.strip():
                        files.append((name, markdown))
                except Exception as e:
                    logger.warning(f"Error processing {name} in ZIP: {e}")
    except Exception as e:
        logger.error(f"Error reading ZIP file: {e}")

    return files
