"""
SAGE-Docs Document Ingestion Pipeline
======================================
Processes uploaded documents (Markdown, HTML, Text, PDF, ZIP) into
searchable chunks indexed in Qdrant vector database.
"""

import os
import io
import re
import hashlib
import logging
import zipfile
from pathlib import Path
from typing import Optional
import yaml

from markdownify import markdownify as md
from bs4 import BeautifulSoup
import subprocess
import shutil
from qdrant_client import QdrantClient
from qdrant_client.http import models
from fastembed import TextEmbedding, SparseTextEmbedding

# Optional imports for DOCX and Excel
try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import openpyxl
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger("SAGE-Ingest")

# Configuration from environment
QDRANT_HOST = os.getenv("QDRANT_HOST", "localhost")
QDRANT_PORT = int(os.getenv("QDRANT_PORT", "6333"))
COLLECTION_NAME = os.getenv("COLLECTION_NAME", "sage_docs")
UPLOAD_DIR = Path(os.getenv("UPLOAD_DIR", "/app/uploads"))

# Chunking settings
CHUNK_SIZE = 1500
CHUNK_OVERLAP = 200

# Embedding settings
EMBEDDING_MODE = os.getenv("EMBEDDING_MODE", "local")
DENSE_MODEL_NAME = os.getenv("DENSE_MODEL_NAME", "sentence-transformers/all-MiniLM-L6-v2")
DENSE_VECTOR_SIZE = int(os.getenv("DENSE_VECTOR_SIZE", "384"))
USE_NOMIC_PREFIX = os.getenv("USE_NOMIC_PREFIX", "false").lower() == "true"

# Global model instances
_dense_model: Optional[TextEmbedding] = None
_sparse_model: Optional[SparseTextEmbedding] = None


def get_dense_model() -> TextEmbedding:
    """Get or create dense embedding model."""
    global _dense_model
    if _dense_model is None:
        logger.info(f"Loading dense model: {DENSE_MODEL_NAME}")
        _dense_model = TextEmbedding(model_name=DENSE_MODEL_NAME)
    return _dense_model


def get_sparse_model() -> SparseTextEmbedding:
    """Get or create sparse BM25 model."""
    global _sparse_model
    if _sparse_model is None:
        logger.info("Loading sparse BM25 model...")
        _sparse_model = SparseTextEmbedding(model_name="Qdrant/bm25")
    return _sparse_model


def get_content_hash(content: str) -> str:
    """Generate MD5 hash of content for deduplication."""
    return hashlib.md5(content.encode()).hexdigest()


def detect_file_type(filename: str, content: bytes) -> str:
    """Detect file type based on extension and content."""
    ext = Path(filename).suffix.lower()
    
    if ext == '.md' or ext == '.markdown':
        return 'markdown'
    elif ext in ['.html', '.htm']:
        return 'html'
    elif ext == '.txt':
        return 'text'
    elif ext == '.pdf':
        return 'pdf'
    elif ext == '.zip':
        return 'zip'
    elif ext == '.docx':
        return 'docx'
    elif ext in ['.xlsx', '.xls']:
        return 'excel'
    elif ext in ['.rst', '.asciidoc', '.adoc']:
        return 'text'  # Treat as plain text
    else:
        # Try to detect from content
        try:
            text = content.decode('utf-8', errors='ignore')[:1000]
            if text.strip().startswith('<!DOCTYPE') or '<html' in text.lower():
                return 'html'
            elif text.startswith('---\n') or re.search(r'^#\s+\w', text, re.MULTILINE):
                return 'markdown'
        except:
            pass
        return 'text'


def convert_html_to_markdown(html_content: str) -> str:
    """Convert HTML to clean Markdown."""
    # Clean up with BeautifulSoup first
    soup = BeautifulSoup(html_content, 'html.parser')
    
    # Remove script and style elements
    for script in soup(["script", "style", "nav", "footer", "header"]):
        script.decompose()
    
    # Convert to markdown
    markdown = md(
        str(soup),
        heading_style="atx",
        code_language_callback=lambda el: el.get('data-language') or el.get('class', [''])[0] if el.get('class') else ''
    )
    
    return markdown.strip()


def extract_pdf_text(pdf_content: bytes) -> str:
    """Extract text from PDF file using olmocr for layout preservation."""
    import tempfile
    
    # olmocr configuration from environment
    olmocr_server = os.getenv("OLMOCR_SERVER", "")  # External vLLM server URL
    olmocr_api_key = os.getenv("OLMOCR_API_KEY", "")  # API key for external providers
    olmocr_model = os.getenv("OLMOCR_MODEL", "allenai/olmOCR-2-7B-1025-FP8")
    
    try:
        # Write PDF to temp file
        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
            tmp.write(pdf_content)
            tmp_path = tmp.name
        
        # Create workspace for olmocr output
        workspace = tempfile.mkdtemp(prefix="olmocr_")
        
        logger.info("Converting PDF with olmocr (this may take a while)...")
        
        # Build command
        cmd = [
            "python", "-m", "olmocr.pipeline", workspace,
            "--markdown", "--pdfs", tmp_path,
            "--model", olmocr_model
        ]
        
        # Add server/API configuration if provided
        if olmocr_server:
            cmd.extend(["--server", olmocr_server])
        if olmocr_api_key:
            cmd.extend(["--api_key", olmocr_api_key])
        
        # Run olmocr pipeline
        result = subprocess.run(
            cmd,
            capture_output=True,
            text=True,
            timeout=600  # 10 minute timeout for large PDFs
        )
        
        if result.returncode != 0:
            logger.error(f"olmocr failed: {result.stderr}")
            return ""
        
        # Read generated markdown
        pdf_stem = Path(tmp_path).stem
        md_file = Path(workspace) / "markdown" / f"{pdf_stem}.md"
        
        if md_file.exists():
            markdown = md_file.read_text()
            logger.info(f"PDF conversion complete: {len(markdown)} chars")
        else:
            logger.warning(f"olmocr did not produce markdown output for {tmp_path}")
            markdown = ""
        
        # Clean up
        os.remove(tmp_path)
        shutil.rmtree(workspace, ignore_errors=True)
        
        return markdown
    except subprocess.TimeoutExpired:
        logger.error("olmocr timed out processing PDF")
        return ""
    except Exception as e:
        logger.error(f"Error extracting PDF with olmocr: {e}")
        return ""


def extract_docx_text(docx_content: bytes) -> str:
    """Extract text from DOCX file."""
    if not DOCX_AVAILABLE:
        logger.warning("python-docx not available, skipping DOCX extraction")
        return ""
    
    try:
        doc = DocxDocument(io.BytesIO(docx_content))
        text_parts = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                # Check if it's a heading
                if para.style.name.startswith('Heading'):
                    level = para.style.name.replace('Heading ', '')
                    try:
                        level = int(level)
                        text_parts.append(f"{'#' * level} {para.text}")
                    except:
                        text_parts.append(f"## {para.text}")
                else:
                    text_parts.append(para.text)
        
        # Also extract text from tables
        for table in doc.tables:
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(' | '.join(cells))
            if rows:
                text_parts.append('\n'.join(rows))
        
        return '\n\n'.join(text_parts)
    except Exception as e:
        logger.error(f"Error extracting DOCX: {e}")
        return ""


def extract_excel_text(excel_content: bytes) -> str:
    """Extract text from Excel file."""
    if not EXCEL_AVAILABLE:
        logger.warning("openpyxl not available, skipping Excel extraction")
        return ""
    
    try:
        wb = openpyxl.load_workbook(io.BytesIO(excel_content), data_only=True)
        text_parts = []
        
        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            text_parts.append(f"## {sheet_name}")
            
            rows = []
            for row in sheet.iter_rows(values_only=True):
                # Convert None to empty string and join
                cells = [str(cell) if cell is not None else '' for cell in row]
                if any(cells):  # Only add non-empty rows
                    rows.append(' | '.join(cells))
            
            if rows:
                text_parts.append('\n'.join(rows))
        
        return '\n\n'.join(text_parts)
    except Exception as e:
        logger.error(f"Error extracting Excel: {e}")
        return ""


def extract_title_from_content(content: str, filename: str) -> str:
    """Extract title from content or use filename."""
    # Try to find markdown header
    match = re.search(r'^#\s+(.+)$', content, re.MULTILINE)
    if match:
        return match.group(1).strip()
    
    # Try YAML frontmatter
    if content.startswith('---'):
        try:
            end = content.find('---', 3)
            if end > 0:
                frontmatter = yaml.safe_load(content[3:end])
                if isinstance(frontmatter, dict) and 'title' in frontmatter:
                    return frontmatter['title']
        except:
            pass
    
    # Use filename without extension
    return Path(filename).stem.replace('_', ' ').replace('-', ' ').title()


def split_text_semantic(text: str, chunk_size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP) -> list[str]:
    """
    Split text into chunks with overlap, respecting code blocks and markdown headers.
    """
    # Protect code blocks
    code_block_pattern = r'```[\s\S]*?```'
    code_blocks = re.findall(code_block_pattern, text)
    placeholders = []
    
    for i, block in enumerate(code_blocks):
        placeholder = f"__CODE_BLOCK_{i}__"
        placeholders.append((placeholder, block))
        text = text.replace(block, placeholder, 1)
    
    # Split by paragraphs first
    paragraphs = re.split(r'\n\n+', text)
    
    chunks = []
    current_chunk = ""
    
    for para in paragraphs:
        # Check if this is a header
        is_header = para.strip().startswith('#')
        
        # If adding this paragraph would exceed chunk size
        if len(current_chunk) + len(para) + 2 > chunk_size:
            if current_chunk:
                chunks.append(current_chunk.strip())
                # Keep overlap
                words = current_chunk.split()
                overlap_words = words[-overlap//10:] if len(words) > overlap//10 else []
                current_chunk = ' '.join(overlap_words) + '\n\n' if overlap_words else ''
        
        # Start new chunk on headers if current chunk is substantial
        if is_header and len(current_chunk) > chunk_size // 3:
            if current_chunk:
                chunks.append(current_chunk.strip())
            current_chunk = para + '\n\n'
        else:
            current_chunk += para + '\n\n'
    
    if current_chunk.strip():
        chunks.append(current_chunk.strip())
    
    # Restore code blocks
    for i, chunk in enumerate(chunks):
        for placeholder, block in placeholders:
            chunk = chunk.replace(placeholder, block)
        chunks[i] = chunk
    
    return chunks


def process_file(
    content: bytes,
    filename: str,
    library: str,
    version: str = "latest"
) -> str:
    """
    Process a single file and return its markdown content.
    """
    file_type = detect_file_type(filename, content)
    
    if file_type == 'markdown':
        return content.decode('utf-8', errors='ignore')
    elif file_type == 'html':
        html = content.decode('utf-8', errors='ignore')
        return convert_html_to_markdown(html)
    elif file_type == 'text':
        return content.decode('utf-8', errors='ignore')
    elif file_type == 'pdf':
        return extract_pdf_text(content)
    elif file_type == 'docx':
        return extract_docx_text(content)
    elif file_type == 'excel':
        return extract_excel_text(content)
    else:
        return content.decode('utf-8', errors='ignore')


def process_zip(
    zip_content: bytes,
    library: str,
    version: str = "latest"
) -> list[tuple[str, str]]:
    """
    Process a ZIP file and return list of (filename, content) tuples.
    """
    files = []
    
    try:
        with zipfile.ZipFile(io.BytesIO(zip_content), 'r') as zf:
            for name in zf.namelist():
                # Skip directories and hidden files
                if name.endswith('/') or '/.' in name or name.startswith('.'):
                    continue
                
                # Skip non-document files
                ext = Path(name).suffix.lower()
                if ext not in ['.md', '.markdown', '.html', '.htm', '.txt', '.pdf', '.rst', '.docx', '.xlsx', '.xls']:
                    continue
                
                try:
                    content = zf.read(name)
                    markdown = process_file(content, name, library, version)
                    if markdown.strip():
                        files.append((name, markdown))
                except Exception as e:
                    logger.warning(f"Error processing {name} in ZIP: {e}")
    except Exception as e:
        logger.error(f"Error reading ZIP file: {e}")
    
    return files


async def ingest_document(
    client: QdrantClient,
    content: bytes,
    filename: str,
    library: str,
    version: str = "latest"
) -> dict:
    """
    Ingest a document into the vector database.
    
    Returns:
        dict with ingestion statistics
    """
    logger.info(f"Ingesting document: {filename} for library {library} v{version}")
    
    # Ensure collection exists
    await ensure_collection(client)
    
    # Detect file type
    file_type = detect_file_type(filename, content)
    
    if file_type == 'zip':
        # Process ZIP archive
        files = process_zip(content, library, version)
        total_chunks = 0
        for fname, markdown in files:
            chunks = await _ingest_markdown(client, markdown, fname, library, version)
            total_chunks += chunks
        return {
            "library": library,
            "version": version,
            "files_processed": len(files),
            "chunks_indexed": total_chunks
        }
    else:
        # Process single file
        markdown = process_file(content, filename, library, version)
        chunks = await _ingest_markdown(client, markdown, filename, library, version)
        return {
            "library": library,
            "version": version,
            "files_processed": 1,
            "chunks_indexed": chunks
        }


async def _ingest_markdown(
    client: QdrantClient,
    markdown: str,
    filename: str,
    library: str,
    version: str
) -> int:
    """Ingest markdown content as chunks."""
    # Extract title
    title = extract_title_from_content(markdown, filename)
    
    # Split into chunks
    chunks = split_text_semantic(markdown)
    
    if not chunks:
        return 0
    
    # Get models
    dense_model = get_dense_model()
    sparse_model = get_sparse_model()
    
    # Save original file
    file_path = save_uploaded_file(markdown.encode(), filename, library, version)
    
    # Generate embeddings
    points = []
    
    for i, chunk in enumerate(chunks):
        # Prepare text for embedding
        embed_text = chunk
        if USE_NOMIC_PREFIX:
            embed_text = f"search_document: {chunk}"
        
        # Generate dense embedding
        dense_embedding = list(dense_model.embed([embed_text]))[0].tolist()
        
        # Generate sparse embedding
        sparse_result = list(sparse_model.embed([chunk]))[0]
        sparse_embedding = models.SparseVector(
            indices=sparse_result.indices.tolist(),
            values=sparse_result.values.tolist()
        )
        
        # Create unique ID
        chunk_id = get_content_hash(f"{library}:{version}:{filename}:{i}:{chunk[:100]}")
        
        point = models.PointStruct(
            id=chunk_id,
            vector={
                "dense": dense_embedding,
                "sparse": sparse_embedding
            },
            payload={
                "content": chunk,
                "library": library,
                "version": version,
                "title": title,
                "file_path": str(file_path),
                "chunk_index": i,
                "total_chunks": len(chunks),
                "type": "document"
            }
        )
        points.append(point)
    
    # Upsert to Qdrant
    if points:
        client.upsert(
            collection_name=COLLECTION_NAME,
            points=points
        )
        logger.info(f"Indexed {len(points)} chunks for {filename}")
    
    return len(points)


def save_uploaded_file(content: bytes, filename: str, library: str, version: str) -> Path:
    """Save uploaded file to disk."""
    # Create directory structure
    save_dir = UPLOAD_DIR / library / version
    save_dir.mkdir(parents=True, exist_ok=True)
    
    # Sanitize filename
    safe_name = re.sub(r'[^\w\-_\.]', '_', filename)
    if not safe_name.endswith('.md'):
        safe_name = Path(safe_name).stem + '.md'
    
    file_path = save_dir / safe_name
    
    with open(file_path, 'wb') as f:
        f.write(content)
    
    return file_path


async def ensure_collection(client: QdrantClient):
    """Ensure the collection exists with proper configuration."""
    collections = client.get_collections().collections
    exists = any(c.name == COLLECTION_NAME for c in collections)
    
    if not exists:
        logger.info(f"Creating collection: {COLLECTION_NAME}")
        client.create_collection(
            collection_name=COLLECTION_NAME,
            vectors_config={
                "dense": models.VectorParams(
                    size=DENSE_VECTOR_SIZE,
                    distance=models.Distance.COSINE
                )
            },
            sparse_vectors_config={
                "sparse": models.SparseVectorParams(
                    index=models.SparseIndexParams(
                        on_disk=False,
                    )
                )
            },
            # INT8 Scalar Quantization for memory reduction
            quantization_config=models.ScalarQuantization(
                scalar=models.ScalarQuantizationConfig(
                    type=models.ScalarType.INT8,
                    always_ram=True
                )
            )
        )
        
        # Create payload indexes for filtering
        client.create_payload_index(
            collection_name=COLLECTION_NAME,
            field_name="library",
            field_schema=models.PayloadSchemaType.KEYWORD
        )
        client.create_payload_index(
            collection_name=COLLECTION_NAME,
            field_name="version",
            field_schema=models.PayloadSchemaType.KEYWORD
        )
        client.create_payload_index(
            collection_name=COLLECTION_NAME,
            field_name="file_path",
            field_schema=models.PayloadSchemaType.KEYWORD
        )
        
        logger.info(f"Collection {COLLECTION_NAME} created successfully")


async def delete_library(client: QdrantClient, library: str, version: str = None) -> int:
    """Delete a library (and optionally specific version) from the index."""
    filter_conditions = [
        models.FieldCondition(
            key="library",
            match=models.MatchValue(value=library)
        )
    ]
    
    if version:
        filter_conditions.append(
            models.FieldCondition(
                key="version",
                match=models.MatchValue(value=version)
            )
        )
    
    # Count before delete
    count_result = client.count(
        collection_name=COLLECTION_NAME,
        count_filter=models.Filter(must=filter_conditions)
    )
    
    # Delete from Qdrant
    client.delete(
        collection_name=COLLECTION_NAME,
        points_selector=models.FilterSelector(
            filter=models.Filter(must=filter_conditions)
        )
    )
    
    # Delete from filesystem
    if version:
        delete_path = UPLOAD_DIR / library / version
    else:
        delete_path = UPLOAD_DIR / library
    
    if delete_path.exists():
        import shutil
        shutil.rmtree(delete_path)
    
    logger.info(f"Deleted {count_result.count} chunks for library {library}" + (f" v{version}" if version else ""))
    
    return count_result.count
